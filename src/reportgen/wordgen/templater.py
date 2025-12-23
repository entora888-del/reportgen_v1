from __future__ import annotations

from typing import TYPE_CHECKING, Callable

from reportgen.config.loader import load_settings
from reportgen.parsers.boring_xml import (
    extract_cover_from_xml,
    extract_report_metadata,
    parse_boring_xml,
)
from reportgen.utils.template_locator import resolve_template_path
from reportgen.wordgen.liquefaction_post import postprocess_liquefaction_block

if TYPE_CHECKING:
    from reportgen.ai import AIOptions


def build_context_from_inputs(
    xml_path: str,
    ai_options: "AIOptions | None" = None,
    log_func: Callable[[str], None] | None = None,
) -> dict:
    cover = extract_cover_from_xml(xml_path)
    metadata = extract_report_metadata(xml_path)
    boring = parse_boring_xml(xml_path)
    settings = load_settings()

    layers = boring.get("layers", [])
    spt_records = boring.get("spt_records", [])
    texts = settings.text_defaults
    defaults = settings.report_defaults

    def _staff_line(role: str, name: str, license_no: str | None = None, include_license: bool = True) -> str:
        if not name:
            return ""
        lic = (license_no or "").strip()
        suffix = f"(地質調査技士第{lic}号)" if include_license and lic else ""
        return f"{role}\u3000{name}{suffix}"

    def _is_sandy(name: str) -> bool:
        return any(tok in name for tok in ("砂", "礫", "砂質", "礫質"))

    def _sand_grade(n: int) -> str:
        if n < 4:
            return "非常に緩い"
        if n < 10:
            return "緩い"
        if n < 30:
            return "中位の"
        if n < 50:
            return "密な"
        return "非常に密な"

    def _clay_grade(n: int) -> str:
        if n < 2:
            return "非常に軟らかい"
        if n < 4:
            return "軟らかい"
        if n < 8:
            return "中位の"
        if n < 15:
            return "硬い"
        if n < 30:
            return "非常に硬い"
        return "固結した"

    def _classify_layer(name: str, values: list[int]) -> tuple[str, str]:
        if not values:
            return "", ""
        sandy = _is_sandy(name)
        grader = _sand_grade if sandy else _clay_grade
        suffix = "相対密度であると評価される。" if sandy else "コンシステンシーであると評価される。"
        min_grade = grader(min(values))
        max_grade = grader(max(values))
        grade_text = min_grade if min_grade == max_grade else f"{min_grade}～{max_grade}"
        return grade_text, suffix

    def _circled_number(idx: int) -> str:
        circled = ["①", "②", "③", "④", "⑤", "⑥", "⑦", "⑧", "⑨", "⑩", "⑪", "⑫", "⑬", "⑭", "⑮", "⑯", "⑰", "⑱", "⑲", "⑳"]
        if 0 <= idx < len(circled):
            return circled[idx]
        return f"({idx + 1})"

    def _format_header(idx: int, name: str, top: float | None, bottom: float | None, thickness: float | None) -> tuple[str, str, str]:
        number = _circled_number(idx)
        if top is None or bottom is None:
            title = name
        else:
            thick = ""
            if thickness is not None:
                thick = f"{{層厚{thickness:.2f}m}}"
            title = f"{name}　GL-{top:.2f}～ {bottom:.2f}m{thick}"
        header_full = f"{number}{title}"
        return number, title, header_full

    def _normalize_sentence(text: str) -> list[str]:
        text = text.replace("¥n", "\n")
        sentences = []
        for line in text.splitlines():
            line = line.strip()
            if not line:
                continue
            if not line.endswith("。"):
                line = f"{line}。"
            sentences.append(line)
        return sentences

    layer_entries: list[dict[str, str]] = []
    for idx, layer in enumerate(layers):
        name = layer.get("name", "")
        top = layer.get("top")
        bottom = layer.get("bottom")
        thickness = layer.get("thickness")
        number, title, header_full = _format_header(idx, name, top, bottom, thickness)
        section_lines = [header_full]

        n_values = [int(v) for v in layer.get("N_values", []) if v is not None]
        n_sentence = ""
        if n_values:
            n_min = min(n_values)
            n_max = max(n_values)
            n_text = str(n_min) if n_min == n_max else f"{n_min}～{n_max}"
            if "盛土" in name:
                n_sentence = f"　標準貫入試験では、N={n_text}を記録する。"
            else:
                grade_text, suffix = _classify_layer(name, n_values)
                if grade_text:
                    n_sentence = f"　標準貫入試験では、N={n_text}を記録し「{grade_text}」{suffix}"
                else:
                    n_sentence = f"　標準貫入試験では、N={n_text}を記録する。"
            if n_sentence:
                section_lines.append(n_sentence)

        obs_sentences = _normalize_sentence(layer.get("observation", ""))
        observation_text = ""
        if obs_sentences:
            first_sentence = obs_sentences[0]
            if not first_sentence.startswith("本層"):
                first_sentence = f"本層は、{first_sentence}"
            observation_text = f"　{first_sentence}"
            for sentence in obs_sentences[1:]:
                observation_text += f"　{sentence}"
            section_lines.append(observation_text)
        else:
            depth_text = ""
            if top is not None and bottom is not None:
                depth_text = f"（GL-{top:.2f}～{bottom:.2f}m）"
            base = f"　本層は、{name}{depth_text}で構成される。"
            if n_sentence:
                base = base.rstrip("。")
                base += n_sentence
            observation_text = base
            section_lines.append(observation_text)

        layer_entries.append(
            {
                "number": number,
                "title": title,
                "header": header_full,
                "observation": observation_text,
                "n_sentence": n_sentence,
            }
        )

    def _project_name() -> str:
        for key in ("survey_name", "cover_title"):
            value = (cover.get(key) or metadata.get(key) or "").strip()
            if value:
                name = value
                for suffix in ("地質調査", "ボーリング調査"):
                    if name.endswith(suffix):
                        name = name[:-len(suffix)]
                        break
                return name
        return metadata.get("borehole_name", "") or ""

    def _format_meter(value: str) -> str:
        try:
            num = float(value)
        except (TypeError, ValueError):
            return value or ""
        return f"{int(round(num))}ｍ"

    def _estimate_spt_tests(length_text: str, layers_info: list[dict]) -> int:
        max_depth = 0.0
        for layer in layers_info or []:
            try:
                bottom = float(layer.get("bottom") or 0.0)
                max_depth = max(max_depth, bottom)
            except (TypeError, ValueError):
                continue
        if not length_text and max_depth > 0:
            length_text = str(max_depth)
        try:
            length = float(length_text)
        except (TypeError, ValueError):
            return 1
        return max(1, int(round(length)) + 1)

    project_name = _project_name()
    borehole_name = metadata.get("borehole_name", "")
    boring_count = defaults.survey_quantity_boring_count or (1 if borehole_name else 0)
    drilling_length = metadata.get("drilling_length", "")
    spt_sites = defaults.survey_quantity_spt_count or (1 if spt_records else 0)
    spt_tests = len(spt_records) if spt_records else _estimate_spt_tests(drilling_length, layers)

    # 調査概要（テンプレ未設定時も手元の情報から生成）
    base_project = project_name or metadata.get("survey_name", "") or ""
    survey_purpose = texts.survey_purpose_template.format(project_name=base_project or "")
    survey_quantity_boring = f"{boring_count}箇所、L＝{_format_meter(drilling_length)}" if boring_count else ""
    survey_quantity_spt = f"{spt_sites}箇所、n＝{spt_tests}回" if spt_sites else ""

    survey_location = metadata.get("survey_location", "")
    if survey_location.endswith("地内") and not survey_location.endswith(" 地内"):
        survey_location = survey_location.replace("地内", " 地内")

    location_phrase = survey_location or metadata.get("survey_location", "") or ""
    if not survey_purpose.strip():
        project_phrase = base_project or "本調査"
        survey_purpose = (
            f"{project_phrase}に関して、{location_phrase or '計画地'}の地盤構成を把握し、"
            "基礎地盤の検討に必要な資料を得ることを目的とする。"
        )

    layers_detail = ""

    if defaults.layer_defaults:
        for idx, override in enumerate(defaults.layer_defaults):
            if idx >= len(layer_entries):
                break
            if not isinstance(override, dict):
                continue
            for key in ("observation", "n_sentence"):
                value = override.get(key) or ""
                if value and not layer_entries[idx].get(key):
                    layer_entries[idx][key] = value

    for entry in layer_entries:
        obs = entry.get("observation", "")
        if obs and not obs.startswith("　"):
            entry["observation"] = "　" + obs.lstrip()
        n_sentence = entry.get("n_sentence", "")
        if n_sentence and not n_sentence.startswith("　"):
            entry["n_sentence"] = "　" + n_sentence.lstrip()

    layer_strings = []
    for entry in layer_entries:
        parts = [entry.get("header", "")]
        if entry.get("observation"):
            parts.append(entry["observation"])
        if entry.get("n_sentence"):
            parts.append(entry["n_sentence"])
        parts = [p for p in parts if p]
        if parts:
            layer_strings.append("\n".join(parts))

    layers_detail = "\n\n".join(layer_strings).strip()

    context = dict(cover)
    context.update({
        "survey_name": metadata.get("survey_name") or cover.get("cover_title", ""),
        "survey_location": survey_location,
        "survey_period_start": metadata.get("survey_period_start", ""),
        "survey_period_end": metadata.get("survey_period_end", ""),
        "survey_purpose": survey_purpose,
        "survey_quantity_boring": survey_quantity_boring,
        "survey_quantity_spt": survey_quantity_spt,
        "survey_quantity_liq": defaults.survey_quantity_liq,
        "survey_company_name": metadata.get("survey_company_name") or cover.get("cover_company_name", ""),
        "survey_company_address": metadata.get("survey_company_address") or settings.company.address,
        "survey_company_tel": metadata.get("survey_company_tel_xml", "") or settings.company.tel,
        "survey_company_fax": settings.company.fax,
        "survey_staff_lead": _staff_line(
            "主任技術者",
            metadata.get("survey_staff_lead_name", ""),
            metadata.get("survey_staff_lead_license", ""),
        ),
        "survey_staff_agent": _staff_line(
            "業務代理人",
            metadata.get("survey_staff_agent_name", ""),
            include_license=False,
        ),
        "survey_staff_field": _staff_line(
            "現場調査員",
            metadata.get("survey_staff_field_name", ""),
            include_license=False,
        ),
        "site_name": boring.get("site_name", ""),
        "groundwater": boring.get("groundwater", ""),
        "layers": layers,
        "layers_detail": layers_detail,
        "layers_rendered": layer_entries,
        "borehole_name": borehole_name,
        "drilling_direction": metadata.get("drilling_direction", "垂直") or "垂直",
        "drilling_length": metadata.get("drilling_length", ""),
        "borehole_elevation": metadata.get("borehole_elevation", ""),
        "groundwater_date_md": metadata.get("groundwater_date_md", ""),
        "natural_drilling_depth": metadata.get("natural_drilling_depth", "") or defaults.natural_drilling_depth,
        "groundwater_depth": metadata.get("groundwater_depth", ""),
        "groundwater_note": metadata.get("groundwater_note", "") or defaults.groundwater_note,
        "area_location_overview": "【記入】調査地周辺の位置・アクセス等を記載してください。",
        "area_hydrology_overview": "【記入】周辺の水文・河川状況を記載してください。",
        "area_geology_overview": "【記入】周辺の地形・地質概要を記載してください。",
        "area_surface_overview": "【記入】表層地質の概要を記載してください。",
        "drilling_summary": texts.drilling_summary,
        "liq_summary_text": texts.liq_summary_text,
        "liq_conclusion_text": texts.liq_conclusion_text,
        "liq_risk_evaluation": texts.liq_risk_evaluation,
        "foundation_consideration_top": texts.foundation_consideration_top,
        "foundation_consideration_bottom": texts.foundation_consideration_bottom,
        "groundwater_overview": texts.groundwater_overview,
        "liq_ground_displacement_150": defaults.liquefaction_ground_displacement[0],
        "liq_ground_displacement_200": defaults.liquefaction_ground_displacement[1],
        "liq_ground_displacement_350": defaults.liquefaction_ground_displacement[2],
        "liq_degree_150": defaults.liquefaction_degree[0],
        "liq_degree_200": defaults.liquefaction_degree[1],
        "liq_degree_350": defaults.liquefaction_degree[2],
        "liq_index_150": defaults.liquefaction_index[0],
        "liq_index_200": defaults.liquefaction_index[1],
        "liq_index_350": defaults.liquefaction_index[2],
        "liq_risk_150": defaults.liquefaction_risk[0],
        "liq_risk_200": defaults.liquefaction_risk[1],
        "liq_risk_350": defaults.liquefaction_risk[2],
    })

    field_line = context.get("survey_staff_field")
    suffix = defaults.field_staff_suffix
    if field_line and suffix:
        addition = suffix
        if not addition.startswith("　"):
            addition = "　" + addition
        if addition not in field_line:
            context["survey_staff_field"] = f"{field_line}{addition}"

    groundwater_status = defaults.groundwater_status
    if groundwater_status:
        context.setdefault("groundwater_status", groundwater_status)

    lead_line = context.get("survey_staff_lead")
    if lead_line and not lead_line.endswith(" "):
        context["survey_staff_lead"] = f"{lead_line} "

    if ai_options and ai_options.enabled:
        from reportgen.ai import AITextEnhancer, AITextGenerationError

        enhancer = AITextEnhancer(ai_options, logger=log_func)
        try:
            ai_updates = enhancer.enhance_context(context, metadata, boring)
        except AITextGenerationError as exc:  # pragma: no cover - passthrough for UI
            raise RuntimeError(str(exc)) from exc
        else:
            context.update(ai_updates)

    return context


def generate_docx_from_template(
    template_path: str | None,
    out_path: str,
    context: dict,
    liq_result: dict | None = None,
    source_xml_path: str | None = None,
):
    from pathlib import Path
    from docxtpl import DocxTemplate
    from reportgen.utils.template_locator import packaged_template_path
    from docx import Document
    from docx.enum.text import WD_BREAK

    has_liq = bool(liq_result and liq_result.get("cases"))

    template_dir_candidates = []
    template_from_xml_dir = False
    if source_xml_path:
        xml_dir = Path(source_xml_path).expanduser().parent
        template_dir_candidates = sorted(xml_dir.glob("*.docx"))

    if template_path:
        tpl_path = resolve_template_path(template_path)
    else:
        if template_dir_candidates:
            tpl_path = template_dir_candidates[0]
            template_from_xml_dir = True
        elif has_liq:
            tpl_path = resolve_template_path(None)
        else:
            # 非液状化では「正解」テンプレがあれば優先、無ければ旧版
            try:
                tpl_path = packaged_template_path("報告書_正解.docx")
            except FileNotFoundError:
                tpl_path = packaged_template_path("報告書_ひな形.docx")

    doc = DocxTemplate(str(tpl_path))
    doc.render(context)
    doc.save(out_path)
    postprocess_liquefaction_block(out_path, liq_result)

    if not has_liq and not template_from_xml_dir:
        _strip_liq_and_debug_sections(out_path)
        _ensure_chapter_page_breaks(out_path)
    else:
        _ensure_chapter_page_breaks(out_path)


def _strip_liq_and_debug_sections(out_path: str):
    """
    液状化セクションとデバッグ出力（配列表記など）を削除する。
    """
    from docx import Document

    try:
        doc = Document(out_path)
    except Exception:
        return

    body = doc.element.body
    def element_text(el) -> str:
        return "".join(el.itertext())

    to_remove = []
    body_started = False
    skip_liq = False
    skip_geo = False
    for el in list(body):
        text = element_text(el)
        text_norm = text.replace(" ", "").replace("　", "")

        # remove debug-like list output
        if text.strip().startswith("[{") or "N_values" in text:
            to_remove.append(el)
            continue

        if "１．調査概要" in text_norm:
            body_started = True

        # remove liquefaction block between 5-1 and 5-2 (body only)
        if body_started and ("５－１" in text_norm or "5－1" in text_norm or "5-1" in text_norm):
            skip_liq = True
            to_remove.append(el)
            continue
        if skip_liq:
            if "５－２" in text_norm or "5－2" in text_norm or "5-2" in text_norm or "支持層" in text_norm:
                skip_liq = False
                # keep this heading
            else:
                to_remove.append(el)
                continue

        # remove geology overview block between 3.* and 4.*
        if body_started and "３．調査地周辺地形地質概要" in text_norm:
            skip_geo = True
            to_remove.append(el)
            continue
        if skip_geo:
            if text_norm.startswith("４．調査結果") or "４．調査結果" in text_norm:
                skip_geo = False
            else:
                to_remove.append(el)
                continue

        # remove any remaining paragraph/table mentioning liquefaction
        if "液状化" in text_norm:
            to_remove.append(el)

    for el in to_remove:
        body.remove(el)

    try:
        doc.save(out_path)
    except Exception:
        pass


def _ensure_chapter_page_breaks(out_path: str):
    """章見出しの前に改ページを挿入（先頭章は除く）"""
    from docx import Document
    from docx.enum.text import WD_BREAK

    try:
        doc = Document(out_path)
    except Exception:
        return

    headings = ("１．", "2．", "３．", "４．", "５．", "1.", "2.", "3.", "4.", "5.")
    first_seen = False
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if any(text.startswith(h) for h in headings):
            if not first_seen:
                first_seen = True
                continue
            run = para.insert_paragraph_before().add_run()
            run.add_break(WD_BREAK.PAGE)
    try:
        doc.save(out_path)
    except Exception:
        pass
